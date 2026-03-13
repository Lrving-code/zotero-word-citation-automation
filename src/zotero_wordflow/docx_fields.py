from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def citation_code(citation_id: str, display_text: str, citation_items: list[dict[str, Any]]) -> str:
    payload = {
        "citationID": citation_id,
        "properties": {"formattedCitation": display_text, "plainCitation": display_text, "noteIndex": 0},
        "citationItems": citation_items,
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    return " ADDIN ZOTERO_ITEM CSL_CITATION " + json.dumps(payload, ensure_ascii=False, separators=(",", ":"))


def build_citation_items(entries: list[dict[str, Any]], zotero_map: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    citation_items: list[dict[str, Any]] = []
    for entry in entries:
        item = zotero_map[entry["cite_key"]]
        citation_item = {
            "id": item["item_id"],
            "uris": [item["uri"]],
            "itemData": {"id": item["item_id"], "type": "article-journal"},
        }
        if entry.get("suppress_author"):
            citation_item["suppressAuthor"] = True
        citation_items.append(citation_item)
    return citation_items


def citation_display_text(segment: dict[str, Any]) -> str:
    prefix = segment.get("prefix", "(")
    suffix = segment.get("suffix", ")")
    body = "; ".join(entry["display_text"] for entry in segment["citations"])
    return f"{prefix}{body}{suffix}"


def add_zotero_field(paragraph, segment: dict[str, Any], zotero_map: dict[str, dict[str, Any]]) -> None:
    citation_id = "_".join(entry["cite_key"] for entry in segment["citations"])
    display_text = citation_display_text(segment)
    instr = citation_code(citation_id, display_text, build_citation_items(segment["citations"], zotero_map))
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)
    for start in range(0, len(instr), 180):
        run_instr = paragraph.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instr[start : start + 180]
        run_instr._r.append(instr_text)
    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)
    paragraph.add_run(display_text)
    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def render_document(manifest: dict[str, Any], import_result: dict[str, Any], output_path: Path) -> None:
    zotero_map = {item["cite_key"]: item for item in import_result["items"]}
    doc = Document()
    document_meta = manifest.get("document", {})
    doc.core_properties.title = document_meta.get("title", "Zotero Wordflow Output")
    for element in document_meta.get("elements", []):
        if element["type"] == "heading":
            paragraph = doc.add_paragraph()
            style = "Title" if str(element.get("level")) == "title" else f"Heading {int(element.get('level', 1))}"
            paragraph.style = style
            paragraph.add_run(element["text"])
            continue
        paragraph = doc.add_paragraph()
        for segment in element.get("segments", []):
            if "text" in segment:
                paragraph.add_run(segment["text"])
                continue
            add_zotero_field(paragraph, segment, zotero_map)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
